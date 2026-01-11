import docx
import re
import os

def is_line_clean(line):
    """
    Tashkhees karein ke line saaf hai ya nahi.
    Agar saaf hai (True) return karein, warna (False).
    """
    line = line.strip()

    # 1. Khali lines ko ignore karein
    if not line:
        return False

    # 2. Ghalat alfaz ya nishanaat (keywords) ko ignore karein
    bad_keywords = [
        '(', ')', '...', 'unreadable', 'unclear', 
        'Image text', 'etc', 'cut off', 'Name cut off'
    ]
    for keyword in bad_keywords:
        if keyword in line:
            return False

    # 3. Numbered list items ko ignore karein (jaise "1.", "123.")
    # Yeh check karta hai ke line number aur period se shuru hoti hai ya nahi
    if re.match(r'^\s*\d+\.', line):
        return False

    # 4. Agar line mein koi angrezi harf (a-z) hai, to usay ignore karein
    if re.search(r'[a-zA-Z]', line):
        return False

    # 5. Agar line sirf numbers hai, to usay ignore karein
    if line.isdigit():
        return False

    # Agar line tamam checks pass kar le, to yeh ek saaf naam hai
    return True

# --- Main Script ---

input_file = 'data.docx'
output_file = 'cleandata.docx'

try:
    # Input file (data.docx) ko load karein
    doc_in = docx.Document(input_file)
    
    # Ek nayi output file (cleandata.docx) banayein
    doc_out = docx.Document()
    
    print(f"'{input_file}' ko parha ja raha hai...")
    
    clean_names_count = 0
    total_lines_processed = 0

    # Input file ke har paragraph (line) ko check karein
    for para in doc_in.paragraphs:
        total_lines_processed += 1
        line_text = para.text
        
        # Check karein ke line saaf hai ya nahi
        if is_line_clean(line_text):
            # Agar saaf hai, to usay nayi file mein add karein
            doc_out.add_paragraph(line_text)
            clean_names_count += 1
        # else:
            # Debugging ke liye: print(f"Ignoring: {line_text}")
            pass

    # Nayi document ko save karein
    doc_out.save(output_file)
    
    print("\nMukammal!")
    print(f"Total {total_lines_processed} lines process ki gayin.")
    print(f"{clean_names_count} saaf naam nikalay gaye.")
    print(f"Saaf data '{output_file}' mein save ho gaya hai.")

except FileNotFoundError:
    print(f"Error: File '{input_file}' nahi mili.")
    print("Bara-e-meharbani, 'data.docx' ko script wali directory mein rakhein.")
except Exception as e:
    print(f"Ek ghalti pesh aayi: {e}")